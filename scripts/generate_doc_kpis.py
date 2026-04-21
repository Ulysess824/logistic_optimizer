import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_kpi_word(output_word_path):
    print(f"Generando documento Word en: {output_word_path}")
    doc = Document()
    
    # Título Principal
    titulo = doc.add_heading("Manual de KPIs: Modelo de Optimización Logística y Financiera", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Este documento elaborado para la presentación del TFM detalla las métricas (KPIs) clave calculadas "
        "por el optimizador, su fundamentación matemática, y la metodología subyacente para demostrar la solidez "
        "del modelo ante el tribunal."
    )
    
    # 1. KPIs Operativos
    doc.add_heading("1. KPIs de Rendimiento Logístico Operativo", level=2)
    
    kpi_op = doc.add_paragraph()
    kpi_op.add_run("Kilómetros Optimizados y Retornos en Vacío: ").bold = True
    kpi_op.add_run("El algoritmo VRPB (Vehicle Routing Problem with Backhauls) calcula la distancia total (Total_km) "
                   "usando el motor geográfico OSRM, considerando la red de carreteras real. Una métrica vital es la "
                   "reducción de 'Kilómetros de Vacío'. Un viaje tradicional (Tratamiento A) envía el camión al cliente "
                   "y lo devuelve vacío al depósito. El optimizador (Tratamiento B) triangula entregas para forzar que el "
                   "camión retorne con material, minimizando drásticamente la tasa de km inútiles.")
    
    doc.add_paragraph(
        "Cálculo de Ahorro en Vacío: Porcentaje de reducción del kilometraje improductivo de la flota tras "
        "la resolución de las matrices de distancia inter-nodo (Plantas y Clientes).", style='List Bullet'
    )
    
    # 2. KPIs Financieros
    doc.add_heading("2. Modelado Financiero y Total Cost of Ownership (TCO)", level=2)
    
    kpi_fin = doc.add_paragraph()
    kpi_fin.add_run("Tasa de Coste TCO Interno (€/km): ").bold = True
    kpi_fin.add_run("A diferencia de modelos básicos basados puramente en coste de combustible, el sistema modela un TCO "
                    "avanzado. La tasa técnica interna (aprox. 1.50 €/km) integra: CAPEX prorrateado mediante la función de Pago (PMT) "
                    "incluyendo Coste Medio Ponderado del Capital (WACC), depreciación, y un horizonte de vida útil (5 años), además "
                    "del OPEX (Combustible, Seguros, Mantenimiento preventivo ponderado por inflación interanual).")
                    
    kpi_fin2 = doc.add_paragraph()
    kpi_fin2.add_run("Backhauling Profit (Evitado de Costo Hundido): ").bold = True
    kpi_fin2.add_run("Si un transportista es forzado a volver sin carga, esa distancia se considera un 'Costo Hundido'. El algoritmo "
                     "monetiza los kilómetros en vacío ahorrados y los incorpora directamente al balance de ingresos operativos. Esto "
                     "se muestra como '+€ Ahorro Sistémico' en el dashboard.")
                     
    kpi_fin3 = doc.add_paragraph()
    kpi_fin3.add_run("Ley de Little y Dimensión de Flota: ").bold = True
    kpi_fin3.add_run("Para evaluar la inversión de CAPEX en la compra de camiones (Asset Heavy vs Asset Light), el motor calcula "
                     "Matemáticamente los vehículos físicos usando la Ley de Little (L = λ × W), cruzando las salidas necesarias "
                     "diarias contra los ciclos logísticos medios de ruta (lead time).")

    # 3. Sostenibilidad y Emisiones
    doc.add_heading("3. KPI de Sostenibilidad y Emisiones GLEC v3", level=2)
    
    kpi_eco = doc.add_paragraph()
    kpi_eco.add_run("CO2 Emissions (FCR Estimator): ").bold = True
    kpi_eco.add_run("Superando las estimaciones volumétricas lineales, el motor implementa las directivas del Global Logistics "
                    "Emissions Council (GLEC). Para vehículos pesados de 40t-44t, se usa una metodología sensible a la carga viva.\n\n"
                    "El cálculo suma un suelo de emisiones en vacío (0.65 kg/km) más un ratio de intensidad (aprox. 17.3 g CO2/t-km) "
                    "que fluctúa dinámicamente según el remanente de pallets a bordo. Así, una ruta de bajada (cargada) emite más "
                    "que su retorno ligero, y un camión eléctrico cuenta puramente por su eficiencia energética (kWh).")

    doc.add_heading("Conclusión para Defensa", level=2)
    doc.add_paragraph("La superioridad de esta simulación recae en que el optimizador no minimiza meras trazas de polígonos, sino "
                      "que minimiza la carga financiera de la organización logística bajo constantes macroeconómicas reales (inflación, WACC).")
                      
    doc.save(output_word_path)
    print("Éxito: Archivo Word documentando KPIs generado con éxito.")


if __name__ == "__main__":
    base = Path(__file__).resolve().parent.parent / "outputs"
    output_path = base / "Documentacion_KPIs_Logistica.docx"
    
    generate_kpi_word(output_path)
