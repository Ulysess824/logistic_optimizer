import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Importar configuración real para consistencia
from logistic_core.config import (
    TARIFA_INTERNA_DIESEL, TARIFA_INTERNA_EV, 
    EXTERNAL_PROVIDER_RATE_PER_KM,
    BASE_TCO_DIESEL, BASE_TCO_EV,
    TCO_HORIZON_YEARS, TCO_WACC, TCO_INFLACION_ANUAL,
    GLEC_EMPTY_FLOOR_KGKM, GLEC_INTENSITY_GTKM,
    SOFTWARE_TMS_CAPEX, DIESEL_CONSUMO_L_100KM, EV_CONSUMO_KWH_KM
)

def add_paragraph(doc, text, style=None, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.runs[0] if p.runs else p.add_run()
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_formula(doc, formula_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n    {formula_text}\n")
    run.font.name = 'Consolas'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(30, 58, 138) # Blue 900
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def generate():
    doc = Document()
    
    # --- TÍTULO ---
    title = doc.add_heading('Anexo Técnico: Metodología de Cálculo y Definición de KPIs', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_paragraph(doc, f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y')}", size=9, italic=True)
    add_paragraph(doc, "Este documento define formalmente la base matemática y conceptual de los indicadores clave de rendimiento (KPIs) presentados en el Dashboard de Optimización Logística para el Trabajo de Fin de Máster.")

    # --- SECCIÓN I ---
    doc.add_heading('I. Indicadores de Eficiencia Operativa', 1)
    
    doc.add_heading('1.1. Milla Vacía Global (%)', 2)
    add_paragraph(doc, "La milla vacía representa la ineficiencia estructural del transporte por carretera. El indicador global se calcula como el sumatorio de kilómetros recorridos sin carga frente al kilometraje total del sistema:")
    add_formula(doc, "MV_Global (%) = [ (Σ Distancia_Vacia) / (Σ Distancia_Total) ] x 100")

    doc.add_heading('1.2. Fill Rate (Llenado)', 2)
    add_paragraph(doc, "Indica la utilización de la capacidad nominal del activo (tráiler estándar de 13.6m).")
    add_formula(doc, "FR (%) = [ (Σ Pallets_Transportados) / 34 ] x 100")

    # --- SECCIÓN II ---
    doc.add_heading('II. Impacto Ambiental (Modelo GLEC v3)', 1)
    add_paragraph(doc, "Las emisiones de CO2 se calculan siguiendo el estándar internacional Global Logistics Emissions Council (GLEC), utilizando factores de intensidad energética para camiones EURO VI de 40 toneladas.")
    
    doc.add_heading('2.1. Cálculo de Emisiones por Trayecto', 2)
    add_paragraph(doc, "La emisión se desglosa por tramo considerando la masa transportada en ese segmento específico:")
    add_formula(doc, "Emisión_CO2 = Distancia x [ Factor_Emision_Base + (Coeficiente_Intensidad x Carga_TN) ]")
    add_paragraph(doc, f"Parámetros: Base = {GLEC_EMPTY_FLOOR_KGKM} kg/km | Intensidad = {GLEC_INTENSITY_GTKM} kg/tn-km.")

    doc.add_heading('2.2. CO2 Evitado', 2)
    add_paragraph(doc, "Representa el ahorro de emisiones respecto al escenario tradicional (inbound/outbound segmentado).")

    # --- SECCIÓN III ---
    doc.add_heading('III. Análisis de Costes y TCO', 1)
    
    doc.add_heading('3.1. Coste Total de Propiedad (TCO) a 5 años', 2)
    add_paragraph(doc, f"El modelo financiero proyecta los costes en un horizonte de {TCO_HORIZON_YEARS} años, integrando variables macroecónomicas y operativas para determinar el coste unitario por kilómetro (€/km).")
    add_formula(doc, "TCO_5años = [ Inversion_Neta + Σ (Coste_Fijo + Coste_Variable)_t / (1 + Tasa_Descuento)^t ] / (Kms_Totales)")
    
    # Normalizar valores de listas para la tabla
    wacc_range = f"{min(TCO_WACC)*100}% - {max(TCO_WACC)*100}%"
    infl_avg = f"{sum(TCO_INFLACION_ANUAL)/len(TCO_INFLACION_ANUAL)*100:.1f}% (Media)"

    add_paragraph(doc, "Parámetros Financieros Base:")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Concepto'
    hdr_cells[1].text = 'Valor de Referencia'
    
    data = [
        ("TCO Objetivo Diésel (€/km)", f"{BASE_TCO_DIESEL} €"),
        ("TCO Objetivo Eléctrico (€/km)", f"{BASE_TCO_EV} €"),
        ("Horizonte Temporal", f"{TCO_HORIZON_YEARS} años"),
        ("Tasa de Descuento (WACC)", wacc_range),
        ("Inflación Proyectada", infl_avg),
        ("Tarifa Mercado (Outsourcing)", f"{EXTERNAL_PROVIDER_RATE_PER_KM} €/km")
    ]
    
    for concept, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = concept
        row_cells[1].text = value

    doc.add_heading('3.2. Modalidades de Inversión y Modelos Matemáticos', 2)
    add_paragraph(doc, "El Coste Total de Propiedad (TCO) se desglosa según la modalidad de adquisición para reflejar con precisión el impacto del flujo de caja:")
    
    add_paragraph(doc, "a) Compra Directa (Inversión propia):", bold=True)
    add_paragraph(doc, "Se prioriza la amortización de la inversión inicial neta. El TCO incluye todos los costes operativos fijos y variables del sistema.")
    add_formula(doc, "TCO_Compra = Inversion_Neta + Σ [ (Coste_Fijo + Coste_Variable)_t / (1 + Tasa_Descuento)^t ]")

    add_paragraph(doc, "b) Leasing Financiero (Financiación):", bold=True)
    add_paragraph(doc, "El coste del activo se distribuye en cuotas que integran principal e intereses. El operador asume el resto de costes operativos (seguros, mantenimiento).")
    add_formula(doc, "TCO_Leasing = Σ [ (Cuota_Leasing_t + Coste_Operativo_t) / (1 + Tasa_Descuento)^t ]")

    add_paragraph(doc, "c) Renting Operativo (Servicio Integral):", bold=True)
    add_paragraph(doc, "Es la modalidad 'Full-Service'. La mayoría de los costes fijos se consolidan en una cuota única de alquiler.")
    add_formula(doc, "TCO_Renting = Σ [ (Cuota_Renting_t + Coste_Energia_t + Coste_Conductor_t) / (1 + Tasa_Descuento)^t ]")

    doc.add_heading('3.3. Ahorro vs. Mercado (Beneficio Sistémico)', 2)
    add_paragraph(doc, "Mide el beneficio económico directo para la compañía al internalizar el servicio operativo mediante rutas circulares optimizadas frente a la contratación de servicios externos a tarifa de mercado (Spot/Contract).")
    add_formula(doc, "Ahorro Sistémico = Σ (Tarifa_Mercado x Dist) - Σ (TCO_Interno x Dist)")

    # --- SECCIÓN IV ---
    doc.add_heading('IV. Metodología de Análisis de Inversión y Eficiencia', 1)
    add_paragraph(doc, "El modelo financiero distingue entre la recuperación de capital inmovilizado y la eficiencia de los servicios de pago por uso.")

    doc.add_heading('4.1. Punto de Equilibrio (Solo Compra Directa)', 2)
    add_paragraph(doc, "Mide el tiempo necesario para amortizar el CAPEX mediante el ahorro operativo bruto (EBITDA).")
    add_formula(doc, "BE_Compra (Años) = CAPEX_Neto / (Ingresos_Internos - OPEX_Operativo)")

    doc.add_heading('4.2. Eficiencia de Servicios (Leasing y Renting)', 2)
    add_paragraph(doc, "Para modalidades sin desembolso inicial masivo, se evalúa la rentabilidad del flujo de caja mensual.")
    add_paragraph(doc, "1. ROI sobre Cuota:", bold=True)
    add_formula(doc, "ROI_Cuota = Ingresos_Anuales / (OPEX_Energía + Cuota_Anual)")
    
    add_paragraph(doc, "2. Ahorro Sistémico vs Mercado:", bold=True)
    add_paragraph(doc, "Diferencia entre el coste de internalización (OPEX + Cuota) y el precio de contratación externa.")
    add_formula(doc, "Ahorro_Anual = (Tarifa_Mercado - TCO_Unitario) x Kilómetros_Totales")

    doc.add_heading('4.3. Capital Liberado', 2)
    add_paragraph(doc, "Ventaja estratégica de mantener la liquidez en balance en lugar de inmovilizarla en activos fijos, valorada al coste de oportunidad del capital (WACC).")

    # --- SECCIÓN V ---
    doc.add_heading('V. Resumen de Precisión y Formateo', 1)
    add_paragraph(doc, "Por directiva de diseño para la defensa del TFM, todos los indicadores financieros han sido normalizados a un solo decimal (X.1) para maximizar la legibilidad y el enfoque en las magnitudes estratégicas.")
    doc.add_page_break()
    doc.add_heading('Conclusión del Análisis', 1)
    add_paragraph(doc, "Los resultados demuestran que la integración de flujos de retorno (Backhauling) permite una dilución de los costes fijos del activo, logrando una reducción del TCO operativo que supera el coste de oportunidad del transporte subcontratado, mejorando simultáneamente la intensidad de carbono del sistema logístico.")

    # GUARDAR
    filename = "Anexo_Tecnico_Metodologia_TFM.docx"
    doc.save(filename)
    print(f"Documento generado: {filename}")

if __name__ == "__main__":
    generate()
